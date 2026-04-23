"""
Document Processing Service
"""

from pathlib import Path
from typing import Tuple
import logging
from pypdf import PdfReader

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Process various document formats"""

    @staticmethod
    def extract_text_from_pdf(file_path: str) -> Tuple[str, int]:
        """
        Extract text and page count from PDF
        Returns: (text, page_count)
        """
        try:
            pdf_reader = PdfReader(file_path)
            text_parts: list[str] = []
            for page in pdf_reader.pages:
                text_parts.append(page.extract_text() or "")

            text = "\n\n".join(text_parts).strip()

            # PyPDF can be weak on slide decks or table-heavy PDFs. If extraction is thin,
            # try PyMuPDF when available before giving up.
            if len(text.split()) < 50:
                try:
                    import fitz

                    doc = fitz.open(file_path)
                    fitz_text = "\n\n".join(page.get_text("text") for page in doc).strip()
                    if len(fitz_text.split()) > len(text.split()):
                        text = fitz_text
                except Exception as fallback_error:
                    logger.warning(f"PyMuPDF fallback unavailable or failed: {fallback_error}")
            
            return text, len(pdf_reader.pages)
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
            raise

    @staticmethod
    def extract_text_from_docx(file_path: str) -> Tuple[str, int]:
        """
        Extract text and page count from DOCX
        """
        try:
            from docx import Document
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text, len(doc.paragraphs)
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            raise

    @staticmethod
    def extract_text_from_txt(file_path: str) -> Tuple[str, int]:
        """
        Extract text from TXT file
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return text, text.count('\n') + 1
        except Exception as e:
            logger.error(f"Error extracting TXT: {e}")
            raise

    @staticmethod
    def process_file(file_path: str) -> Tuple[str, int]:
        """
        Process file based on extension
        Returns: (extracted_text, page_count)
        """
        path = Path(file_path)
        extension = path.suffix.lower()

        if extension == '.pdf':
            return DocumentProcessor.extract_text_from_pdf(file_path)
        elif extension == '.docx':
            return DocumentProcessor.extract_text_from_docx(file_path)
        elif extension == '.txt':
            return DocumentProcessor.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")

document_processor = DocumentProcessor()
